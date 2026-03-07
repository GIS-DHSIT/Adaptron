import pytest

from adaptron.core.registry import global_registry
from adaptron.ingest.docx import DOCXIngester
from adaptron.ingest.models import DataSource, SourceType


def test_registered_as_ingester_docx():
    cls = global_registry.get("ingester", "docx")
    assert cls is DOCXIngester


def test_ingest_docx_file(tmp_path):
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    docx_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Hello Adaptron")
    doc.add_paragraph("This is a test document.")
    doc.save(str(docx_path))

    ingester = DOCXIngester()
    source = DataSource(source_type=SourceType.DOCX, path=str(docx_path))
    docs = ingester.ingest(source)

    assert len(docs) == 1
    assert "Hello Adaptron" in docs[0].content
    assert "test document" in docs[0].content
    assert docs[0].metadata["paragraph_count"] == 2
    assert docs[0].source_ref == str(docx_path)


def test_ingest_missing_file():
    ingester = DOCXIngester()
    source = DataSource(source_type=SourceType.DOCX, path="/nonexistent/file.docx")
    with pytest.raises(FileNotFoundError):
        ingester.ingest(source)
