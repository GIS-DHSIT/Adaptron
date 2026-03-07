from __future__ import annotations

from pathlib import Path

from adaptron.core.registry import register_plugin
from adaptron.ingest.base import BaseIngester
from adaptron.ingest.models import DataSource, RawDocument


@register_plugin("ingester", "docx")
class DOCXIngester(BaseIngester):
    def ingest(self, source: DataSource) -> list[RawDocument]:
        path = Path(source.path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {path}")
        from docx import Document

        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        content = "\n".join(paragraphs)
        return [
            RawDocument(
                content=content,
                metadata={
                    "paragraph_count": len(paragraphs),
                },
                source_ref=str(path),
            )
        ]

    def supported_types(self) -> list[str]:
        return ["docx"]
